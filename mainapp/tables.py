from docx.shared import Cm, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
import docx
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

#
# ?????????????? ?????? ???????????? ??????
#
table_reports = (Inches(4), Inches(3), Inches(1), Inches(3), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5))  #???????????? ???????????? ???????????????? ?????????????? ???????????? ????????????????

def add_table_reports(doc): # ?????????????? ?????????????????? ??????????????
    table = doc.add_table(rows=2, cols=8)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '????????????'
    hdr_cells[1].text = '?????????????????????? ???? ?????? 61850'
    hdr_cells[2].text = '???????????????? ????????????????'
    hdr_cells[4].text = '???????????? ????????????.'
    hdr_cells[5].text = '???????????????? ???? ?????????????? ??????????????'
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_vertical_cell_direction(hdr_cells[4], 'btLr') # ?? 5 ?????????????? ?????????????????? ???????????????????????? ??????????????????????

    set_repeat_table_header(table.rows[0]) # ???????????????????? ?????????????????? ???? ???????? ????????????????

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    hdr_cells = table.rows[1].cells # ???????????? ???????????? ?????????????????? ??????????????
    hdr_cells[2].text = '????????????????????????'
    hdr_cells[3].text = '????????????'
    hdr_cells[5].text = '?? ?? ??'
    hdr_cells[6].text = '?? ?? ??'
    hdr_cells[7].text = '?? ?? ??'

    set_repeat_table_header(table.rows[1])  # ???????????????????? ?????????????????? ???? ???????? ????????????????
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_border(hdr_cells[i], bottom={"val": "double"}) # ???????????????????????? ?????????????????? ?????????????? ????????????


    # ?????????????????? ?????????????????? ?????????????? ?????????????????? ??????????
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(0, 6))
    table.cell(0, 7).merge(table.cell(1, 7))

    table.style = '?????????? ??????????????51'
    table.allow_autofit = False

    for row in table.rows:
        for idx, width in enumerate(table_reports):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # ?????????????????? ???????????? ??????????????, ?????????? ?????????????? ?????????????????????????? ??????????????????
    return table

def add_row_table_reports(table, tuple2Add):  # ?????????????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    leng=len(table.rows)
    if (leng==3):
        # ???????? ?????? ???????????? ??????????????, ???? ???????????? ???????????? ?????????????? ??????????
        row_cells = table.rows[2].cells
        for i in range(0, 8):
            set_cell_border(row_cells[i], top={"val": "double"})

    for idx in range(0, 8):
        row.cells[idx].text = str(tuple2Add[idx])
        row.cells[idx].width = table_reports[idx]
        set_cell_vertical_alignment(row.cells[idx], align="center")
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ?????????? ????????????????????????'
    row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[2].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[3].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[4].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[5].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[6].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[7].paragraphs[0].style = '?????? ?????????????? ??????????'

    return table

def add_spec_row_table_reports(table, tuple2Add):  # ?????????????????? ???????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    leng = len(table.rows)
    if (leng == 3):
        # ???????? ?????? ???????????? ??????????????, ???? ???????????? ???????????? ?????????????? ??????????
        row_cells = table.rows[2].cells
        for i in range(0, 8):
            set_cell_border(row_cells[i], top={"val": "double"})
    row.cells[0].text = str(tuple2Add[0])
    row.cells[0].width = table_reports[0]
    set_cell_vertical_alignment(row.cells[0], align="center")
    row.cells[1].text = str(tuple2Add[1])
    row.cells[1].width = table_reports[1]
    set_cell_vertical_alignment(row.cells[1], align="center")

    for idx in range(0, 8):
        row.cells[idx].paragraphs[0].style = '?????? ?????????????? ??????????'

    table.cell(leng-1, 1).merge(table.cell(leng-1, 7))
    # ?????????????? ???????????? ???????????? - ?????? ?????? ?????? ???????????? ????????????
    shading_elm1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.cell(leng - 1, 0)._tc.get_or_add_tcPr().append(shading_elm1)
    shading_elm2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.cell(leng - 1, 1)._tc.get_or_add_tcPr().append(shading_elm2)


#
# ?????????????? ?????? ???????????? ???? ????????????????
#

# ???? ?????????????????????? ????????????????????????????
table_sg_sw = (Inches(2), Inches(7), Inches(3))  #???????????? ???????????? ????????????????
def add_table_sg_sw(doc): # ?????????????? ?????????????????????? ????????????????????????????
    table = doc.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '????????????????????????'
    hdr_cells[1].text = '??????????????????'
    hdr_cells[2].text = '?????????????????? ?????????????????????? ????????????????????????????'

    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # ???????????????????? ?????????????????? ???? ???????? ????????????????

    table.style = '?????????? ??????????????51'
    table.allow_autofit = False

    for row in table.rows:
        for idx, width in enumerate(table_sg_sw):
            row.cells[idx].width = width
    return table

def add_row_table_sg_sw(table, tuple2Add):  # ?????????????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    #leng=len(table.rows)
    print('tuple=========', tuple2Add)
    for idx in range(0, 3):
        row.cells[idx].text = str(tuple2Add[idx])
        #row.cells[idx].width = table_sg_sw[idx]
        set_cell_vertical_alignment(row.cells[idx], align="center")

    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[2].paragraphs[0].style = '?????? ?????????????? ??????????'

    return table

def add_row_table_sg_sw_empty(table, tuple2Add):  # ?????????????? ???? ?????????????? ?????????????? add_row_table_sg_sw - ???? ???????????? ???????????? ?????????????? ?????? ???????????????????????? ??????????
    row = table.add_row()
    for idx in range(0, 3):
        row.cells[idx].text = str(tuple2Add[idx])
        #row.cells[idx].width = table_sg_sw[idx]
        set_cell_vertical_alignment(row.cells[idx], align="center")
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ??????????' # ?????? ?????????? ?????????????? !!!!!
    row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[2].paragraphs[0].style = '?????? ?????????????? ??????????'
    return table

def merge_table_sg_sw(table, act_row, count):
    table.cell(act_row-1, 0).merge(table.cell(act_row-count, 0))
    text = table.cell(act_row-1, 0).text.replace('\n','')
    table.cell(act_row - 1, 0).text = text.strip()
    row = table.rows[act_row - 1]
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ??????????'
    table.cell(act_row - 1, 1).merge(table.cell(act_row - count, 1))
    text2 = table.cell(act_row - 1, 1).text.replace('\n','')
    table.cell(act_row - 1, 1).text = text2.strip()
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    return table

def merge_table_sg_sw_header(table):
    num_row = len(table.rows)
    table.cell(num_row-1, 0).merge(table.cell(num_row-1, 2))
    text = table.cell(num_row-1, 0).text.replace('\n','')
    table.cell(num_row - 1, 0).text = text.strip()
    # ?????????? ?????????????? ?????????????????????? ???????????????? ?????????? ????????????, ?????????????????????? ?????? ????????????
    row = table.rows[num_row-1]
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ?????????? ??????????????'
    return table

def add_row_table_sg_sw_final(table): # ?????????????????? ?????????????????? ?????????????? ?? ??????????????
    row = table.add_row()
    table.cell(len(table.rows)-1, 0).merge(table.cell(len(table.rows)-1, 2))
    row.cells[0].text = '* - ???????????????? ???????????????????????? ?????????????????????????? ???? ??????????????????'
    set_cell_vertical_alignment(row.cells[0], align="center")
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ??????????' # ?????? ?????????? ?????????????? !!!!!
    return table


#
# ?????????????? ?????? ???????????? ???? ???????????????? GOOSE
#
table_inputs = (Inches(2), Inches(7), Inches(3))  #???????????? ???????????? ????????????????
def add_table_inputs(doc): # ?????????????? ?????????????????????? ????????????????????????????
    table = doc.add_table(rows=1, cols=3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '????????'
    hdr_cells[1].text = '????????????????'
    hdr_cells[2].text = '????????????????????'

    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # ???????????????????? ?????????????????? ???? ???????? ????????????????

    table.style = '?????????? ??????????????51'
    table.allow_autofit = False

    for row in table.rows:
        for idx, width in enumerate(table_inputs):
            row.cells[idx].width = width
    return table

def add_row_table_inputs(table, tuple2Add):  # ?????????????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    for idx in range(0, 3):
        row.cells[idx].text = str(tuple2Add[idx])
        set_cell_vertical_alignment(row.cells[idx], align="center")

    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[2].paragraphs[0].style = '?????? ?????????????? ??????????'

    return table



#
# ?????????? ?????????????? MMS ??????
#
table_reports_new = (Inches(4), Inches(3), Inches(1), Inches(0.5), Inches(0.5), Inches(0.5), Inches(2), Inches(1))  #???????????? ???????????? ???????????????? ?????????????? ???????????? ????????????????

def add_table_reports_new(doc): # ?????????? ?????????????? ?????????????????? ??????????????
    table = doc.add_table(rows=2, cols=8)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '???????????????????????????? ????????????'
    hdr_cells[2].text = '????'
    hdr_cells[3].text = '??????'
    hdr_cells[4].text = '??????'
    hdr_cells[5].text = '??????'
    hdr_cells[6].text = '?????????????????????? ?????????????? ?? ??????'
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_vertical_cell_direction(hdr_cells[3], 'btLr') # ?? 4 ?????????????? ?????????????????? ???????????????????????? ??????????????????????
    set_vertical_cell_direction(hdr_cells[4], 'btLr') # ?? 5 ?????????????? ?????????????????? ???????????????????????? ??????????????????????
    set_vertical_cell_direction(hdr_cells[5], 'btLr') # ?? 6 ?????????????? ?????????????????? ???????????????????????? ??????????????????????

    set_repeat_table_header(table.rows[0]) # ???????????????????? ?????????????????? ???? ???????? ????????????????

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    hdr_cells = table.rows[1].cells # ???????????? ???????????? ?????????????????? ??????????????
    hdr_cells[0].text = '????????????????????????'
    hdr_cells[1].text = '????????????'
    hdr_cells[6].text = 'DO'
    hdr_cells[7].text = 'DA'

    set_repeat_table_header(table.rows[1])  # ???????????????????? ?????????????????? ???? ???????? ????????????????
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = '?????? ?????????????? ??????????????????'
        set_cell_border(hdr_cells[i], bottom={"val": "double"}) # ???????????????????????? ?????????????????? ?????????????? ????????????


    # ?????????????????? ?????????????????? ?????????????? ?????????????????? ??????????
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(0, 7))

    table.style = '?????????? ??????????????51'
    table.allow_autofit = False

    for row in table.rows:
        for idx, width in enumerate(table_reports_new):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # ?????????????????? ???????????? ??????????????, ?????????? ?????????????? ?????????????????????????? ??????????????????
    return table

def add_row_table_reports_new(table, tuple2Add):  # ?????????????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    leng=len(table.rows)
    if (leng==3):
        # ???????? ?????? ???????????? ??????????????, ???? ???????????? ???????????? ?????????????? ??????????
        row_cells = table.rows[2].cells
        for i in range(0, 8):
            set_cell_border(row_cells[i], top={"val": "double"})

    for idx in range(0, 8):
        row.cells[idx].text = str(tuple2Add[idx])
        row.cells[idx].width = table_reports[idx]
        set_cell_vertical_alignment(row.cells[idx], align="center")
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[0].paragraphs[0].style = '?????? ?????????????? ?????????? ????????????????????????'
    row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[1].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[2].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[3].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[4].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[5].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row.cells[6].paragraphs[0].style = '?????? ?????????????? ??????????'
    row.cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row.cells[7].paragraphs[0].style = '?????? ?????????????? ??????????'

    return table

def add_spec_row_table_reports_new(table, tuple2Add):  # ?????????????????? ???????????? ???????????? ???? ???????????????????? ?? ?????????????? ???????????????? ????????????????
    row = table.add_row()
    leng = len(table.rows)
    if (leng == 3):
        # ???????? ?????? ???????????? ??????????????, ???? ???????????? ???????????? ?????????????? ??????????
        row_cells = table.rows[2].cells
        for i in range(0, 8):
            set_cell_border(row_cells[i], top={"val": "double"})
    row.cells[0].text = str(tuple2Add[0])
    row.cells[0].width = table_reports[0]
    set_cell_vertical_alignment(row.cells[0], align="center")
    row.cells[1].text = str(tuple2Add[1])
    row.cells[1].width = table_reports[1]
    set_cell_vertical_alignment(row.cells[1], align="center")

    for idx in range(0, 8):
        row.cells[idx].paragraphs[0].style = '?????? ?????????????? ??????????'

    table.cell(leng-1, 1).merge(table.cell(leng-1, 7))
    # ?????????????? ???????????? ???????????? - ?????? ?????? ?????? ???????????? ????????????
    shading_elm1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.cell(leng - 1, 0)._tc.get_or_add_tcPr().append(shading_elm1)
    shading_elm2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.cell(leng - 1, 1)._tc.get_or_add_tcPr().append(shading_elm2)
